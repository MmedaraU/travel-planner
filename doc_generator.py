from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import os
import database as db


# ===================================================================
#  HELPER: FORMAT DATE TO DD-MM-YYYY
# ===================================================================
def format_date_doc(date_str, output_format="%d-%m-%Y"):
    """Convert ISO date string to DD-MM-YYYY for documents."""
    if not date_str:
        return ""
    try:
        dt = datetime.fromisoformat(date_str)
        return dt.strftime(output_format)
    except:
        return date_str


def format_datetime_doc(dt_str, output_format="%d-%m-%Y %H:%M"):
    """Convert ISO datetime string to DD-MM-YYYY HH:MM."""
    if not dt_str:
        return ""
    try:
        dt = datetime.fromisoformat(dt_str)
        return dt.strftime(output_format)
    except:
        return dt_str


def get_currency_symbol(currency):
    """Return the symbol for a given currency code."""
    symbols = {
        "USD": "$",
        "EUR": "€",
        "GBP": "£",
        "NGN": "₦",
        "JPY": "¥",
        "BRL": "R$",
        "CAD": "C$",
        "AUD": "A$",
        "CHF": "Fr",
        "CNY": "¥",
        "INR": "₹",
    }
    return symbols.get(currency, "$")


# ===================================================================
#  1.  ITINERARY DOCUMENT (uses per‑trip Display Currency)
# ===================================================================
def generate_itinerary_doc(
    executive,
    items,
    stops,
    departure_city,
    departure_region,
    departure_country,
    trip_id,
    trip_budget,
    currency_symbol="$",
    currency_code="USD",
    base_currency=None,
    convert_to_base=False,
):
    """
    Generate a Word itinerary.

    Parameters:
        currency_symbol: Symbol of the trip's Display Currency (e.g., "$").
        currency_code: Code of the trip's Display Currency (e.g., "USD").
        base_currency: Optional base currency for conversion (ignored if convert_to_base=False).
        convert_to_base: If True, converts all costs to base_currency using snapshot rates.
                         If False, displays costs in the original currency with the trip's
                         Display Currency symbol.
    """
    doc = Document()

    # --- Title ---
    title = doc.add_heading(f'Itinerary for {executive["name"]}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Build Departure Display ---
    dep_parts = [p for p in [departure_city, departure_region, departure_country] if p]
    departure_display = ", ".join(dep_parts) if dep_parts else ""

    # --- Route (Departure → Stops) ---
    if stops:
        stop_names = []
        for stop in stops:
            name = stop["city"]
            location_parts = []
            if stop.get("region"):
                location_parts.append(stop["region"])
            if stop.get("country"):
                location_parts.append(stop["country"])
            if location_parts:
                name += f" ({', '.join(location_parts)})"
            stop_names.append(name)

        if departure_display:
            route = f"📍 {departure_display} → " + " → ".join(stop_names)
        else:
            route = " → ".join(stop_names)
        sub = doc.add_paragraph(f"Trip Route: {route}")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Show dates for each stop with full location
        stop_dates = []
        for stop in stops:
            loc = stop["city"]
            location_parts = []
            if stop.get("region"):
                location_parts.append(stop["region"])
            if stop.get("country"):
                location_parts.append(stop["country"])
            if location_parts:
                loc += f" ({', '.join(location_parts)})"
            start = format_date_doc(stop["start_date"])
            end = format_date_doc(stop["end_date"])
            stop_dates.append(f"{loc}: {start} - {end}")
        if stop_dates:
            doc.add_paragraph(" | ".join(stop_dates)).alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )
    else:
        sub = doc.add_paragraph(f'Destination: {executive.get("destination", "N/A")}')
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # --- Executive Profile Summary ---
    doc.add_heading("Executive Profile", level=1)
    doc.add_paragraph(f"Seat: {executive.get('seat_preference', 'Not set')}")
    doc.add_paragraph(f"Hotel Loyalty: {executive.get('hotel_loyalty', 'Not set')}")
    doc.add_paragraph(f"Dietary: {executive.get('dietary_restrictions', 'None')}")
    doc.add_paragraph(f"Meal Preference: {executive.get('meal_preference', 'Not set')}")
    doc.add_paragraph()

    # --- Daily Agenda ---
    doc.add_heading("Daily Agenda", level=1)

    # Determine which currency and symbol to use for display
    if convert_to_base and base_currency:
        display_symbol = get_currency_symbol(base_currency)
    else:
        display_symbol = currency_symbol

    for item in items:
        start_display = format_datetime_doc(item["datetime_start"])
        end_time = (
            datetime.fromisoformat(item["datetime_end"]).strftime("%H:%M")
            if item["datetime_end"]
            else "TBD"
        )
        status_icon = "✅" if item.get("is_confirmed") else "📌"

        # Compute cost
        orig_cost = item.get("cost", 0)
        if convert_to_base and base_currency:
            snapshot_rate = item.get("exchange_rate_snapshot", 1.0)
            cost_str = f"{display_symbol}{orig_cost * snapshot_rate:.2f}"
            # Optionally include original currency info
            orig_currency = item.get("cost_currency", "")
            if orig_currency and orig_currency != base_currency:
                cost_str += f" ({orig_cost:.2f} {orig_currency})"
        else:
            cost_str = f"{display_symbol}{orig_cost:.2f}" if orig_cost else ""

        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{start_display} – {end_time}  |  ").bold = True
        p.add_run(f"{status_icon} {item['description']} ({item['item_type']})")
        if cost_str:
            p.add_run(f"  |  Cost: {cost_str}")
        if item.get("confirmation_code"):
            p.add_run(f"  |  Conf: {item['confirmation_code']}")

    # --- Conflict Warnings ---
    from utils import detect_conflicts

    conflicts = detect_conflicts(items)
    if conflicts:
        doc.add_heading("⚠️ Conflicts Detected", level=1)
        for c in conflicts:
            doc.add_paragraph(c, style="List Bullet")

    # --- Spending Summary (converted to base if requested) ---
    doc.add_page_break()
    doc.add_heading("💰 Trip Spending Summary", level=1)

    # Compute converted totals
    total_estimated = 0
    total_confirmed = 0
    total_all = 0
    for item in items:
        cost = item.get("cost", 0)
        if convert_to_base and base_currency:
            snapshot_rate = item.get("exchange_rate_snapshot", 1.0)
            conv_cost = cost * snapshot_rate
        else:
            conv_cost = cost
        total_all += conv_cost
        if item.get("is_confirmed"):
            total_confirmed += conv_cost
        else:
            total_estimated += conv_cost

    table = doc.add_table(rows=3, cols=2)
    table.style = "Light Grid Accent 1"
    table.cell(0, 0).text = "Category"
    table.cell(0, 1).text = "Amount"
    table.cell(1, 0).text = "Total Estimated (Quoted)"
    table.cell(1, 1).text = f"{display_symbol}{total_estimated:.2f}"
    table.cell(2, 0).text = "Total Confirmed (Booked)"
    table.cell(2, 1).text = f"{display_symbol}{total_confirmed:.2f}"

    doc.add_paragraph(f"\nTotal Trip Spend: {display_symbol}{total_all:.2f}")
    if trip_budget > 0:
        remaining = trip_budget - total_all
        doc.add_paragraph(
            f"Budget: {display_symbol}{trip_budget:.2f}  |  Remaining: {display_symbol}{remaining:.2f}"
        )

    # Footnote about conversion if used
    if convert_to_base and base_currency:
        doc.add_paragraph(
            f"Note: Costs converted to {base_currency} using the exchange rate at the time of each expense (snapshot).",
            style="Normal",
        )

    # --- Footer ---
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(
        f'Generated on {datetime.now().strftime("%d-%m-%Y at %I:%M %p")}  |  Currency: {base_currency if convert_to_base and base_currency else currency_code}'
    )
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True

    # --- Save and return ---
    os.makedirs("generated_itineraries", exist_ok=True)
    filename = f"generated_itineraries/{executive['name']}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    doc.save(filename)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream, filename


# ===================================================================
#  2.  EXECUTIVE PROFILE DOCUMENT (no currency changes)
# ===================================================================
def generate_executive_profile_doc(profile_data, exec_id, currency_symbol="$"):
    """Generate a Word profile with all executive details + memberships."""
    doc = Document()

    # Title
    title = doc.add_heading(f'Executive Profile: {profile_data["Name"]}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f'Company: {profile_data["Company"]}')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Preference table
    table = doc.add_table(rows=10, cols=2)
    table.style = "Light Grid Accent 1"

    rows_data = [
        ("📧 Email", profile_data.get("Email", "Not set")),
        ("🕐 Timezone", profile_data.get("Timezone", "Not set")),
        ("💺 Seat Preference", profile_data.get("Seat Preference", "Not set")),
        ("✈️ Preferred Airline", profile_data.get("Preferred Airline", "Not set")),
        ("🛂 Passport Number", profile_data.get("Passport Number", "Not set")),
        ("✅ TSA PreCheck", profile_data.get("TSA PreCheck", "Not set")),
        ("🥗 Meal Preference", profile_data.get("Meal Preference", "Not set")),
        ("🏨 Hotel Loyalty", profile_data.get("Hotel Loyalty", "Not set")),
        ("✈️ Frequent Flyer #", profile_data.get("Frequent Flyer", "Not set")),
        ("🥗 Dietary Restrictions", profile_data.get("Dietary", "None")),
    ]

    for i, (label, value) in enumerate(rows_data):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value

    doc.add_paragraph()

    # --- Memberships ---
    memberships = db.get_memberships(exec_id)
    if memberships:
        doc.add_heading("✈️ Memberships", level=1)

        airline_mems = [m for m in memberships if m["category"] == "airline"]
        hotel_mems = [m for m in memberships if m["category"] == "hotel"]
        car_mems = [m for m in memberships if m["category"] == "car"]

        if airline_mems:
            doc.add_paragraph("**Airlines:**", style="List Bullet")
            for m in airline_mems:
                doc.add_paragraph(
                    f"{m['program_name']}: {m['membership_number']}",
                    style="List Bullet 2",
                )

        if hotel_mems:
            doc.add_paragraph("**Hotels:**", style="List Bullet")
            for m in hotel_mems:
                doc.add_paragraph(
                    f"{m['program_name']}: {m['membership_number']}",
                    style="List Bullet 2",
                )

        if car_mems:
            doc.add_paragraph("**Car Rentals:**", style="List Bullet")
            for m in car_mems:
                doc.add_paragraph(
                    f"{m['program_name']}: {m['membership_number']}",
                    style="List Bullet 2",
                )

    # Company & Finance
    doc.add_heading("Company & Finance Details", level=1)
    doc.add_paragraph(f"Cost Center: {profile_data.get('Cost Center', 'Not set')}")
    doc.add_paragraph(f"Policy Notes: {profile_data.get('Policy Notes', 'None')}")

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(
        f'Generated on {datetime.now().strftime("%d-%m-%Y at %I:%M %p")}'
    )
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


# ===================================================================
#  3.  SPENDING REPORT (aggregate)
# ===================================================================
def generate_spending_report_doc(
    filter_name,
    summary_data,
    start_date,
    end_date,
    currency_symbol="$",
    base_currency="USD",
):
    """
    Generate an aggregate spending report.
    Note: Totals are shown in base_currency but are NOT automatically converted
    from the raw sums. For accurate conversion per trip, use per‑trip data.
    """
    doc = Document()

    title = doc.add_heading("Executive Travel Spending Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"Executive: {filter_name}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Date Range: {start_date or 'All'} to {end_date or 'All'}")
    doc.add_paragraph()

    # Aggregates (these are raw sums – display in base_currency symbol)
    total_budget = sum(t["budget"] for t in summary_data)
    total_spent = sum(t["total_spent"] for t in summary_data)
    total_confirmed = sum(t["confirmed_spent"] for t in summary_data)
    total_estimated = sum(t["estimated_spent"] for t in summary_data)

    symbol = get_currency_symbol(base_currency)

    doc.add_heading("Aggregate Summary", level=1)
    doc.add_paragraph(f"Total Trips: {len(summary_data)}")
    doc.add_paragraph(f"Total Budget: {symbol}{total_budget:,.2f}")
    doc.add_paragraph(f"Total Spent: {symbol}{total_spent:,.2f}")
    doc.add_paragraph(f"Total Confirmed (Booked): {symbol}{total_confirmed:,.2f}")
    doc.add_paragraph(f"Total Estimated (Quoted): {symbol}{total_estimated:,.2f}")
    doc.add_paragraph()

    # Trip‑level breakdown
    doc.add_heading("Trip‑Level Breakdown", level=1)
    table = doc.add_table(rows=1, cols=7)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    headers = [
        "Executive",
        "Company",
        "Destination",
        "Budget",
        "Total Spent",
        "Confirmed",
        "Status",
    ]
    for i, h in enumerate(headers):
        hdr[i].text = h

    for trip in summary_data:
        row_cells = table.add_row().cells
        row_cells[0].text = trip["executive_name"]
        row_cells[1].text = trip["company_name"]
        row_cells[2].text = trip["destination"]
        row_cells[3].text = f"{symbol}{trip['budget']:.2f}"
        row_cells[4].text = f"{symbol}{trip['total_spent']:.2f}"
        row_cells[5].text = f"{symbol}{trip['confirmed_spent']:.2f}"
        row_cells[6].text = trip["status"]

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(
        f'Generated on {datetime.now().strftime("%d-%m-%Y at %I:%M %p")}  |  Base Currency: {base_currency}'
    )
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


# ===================================================================
#  4.  EXPENSE REPORT (daily breakdown with receipt thumbnails)
#      Always converts to base_currency using snapshot rates
# ===================================================================
def generate_expense_report_doc(
    executive,
    items,
    stops,
    departure_city,
    departure_region,
    departure_country,
    trip_id,
    trip_budget,
    trip_name,
    currency_symbol="$",
    base_currency="USD",
):
    """
    Generate a detailed expense report with:
    - Departure (home base) with structured city, region, country
    - Multi‑city stops (each with city, region, country)
    - Daily grouped items with DD-MM-YYYY
    - Receipt images embedded (thumbnails)
    - Daily subtotals and final summary
    - All costs converted to base_currency using snapshot rates

    Note: currency_symbol is kept for compatibility but the symbol used is derived from base_currency.
    """
    doc = Document()

    # --- Build Departure Display ---
    dep_parts = [p for p in [departure_city, departure_region, departure_country] if p]
    departure_display = ", ".join(dep_parts) if dep_parts else ""

    # --- Header ---
    title = doc.add_heading("Executive Expense Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(
        f"{executive['name']}  |  {executive.get('company_name', '')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if stops:
        stop_names = []
        for stop in stops:
            name = stop["city"]
            location_parts = []
            if stop.get("region"):
                location_parts.append(stop["region"])
            if stop.get("country"):
                location_parts.append(stop["country"])
            if location_parts:
                name += f" ({', '.join(location_parts)})"
            stop_names.append(name)

        if departure_display:
            route = f"📍 {departure_display} → " + " → ".join(stop_names)
        else:
            route = " → ".join(stop_names)
        doc.add_paragraph(f"Trip: {trip_name}  |  Route: {route}").alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        # Stop dates with full location
        stop_info = []
        for stop in stops:
            loc = stop["city"]
            location_parts = []
            if stop.get("region"):
                location_parts.append(stop["region"])
            if stop.get("country"):
                location_parts.append(stop["country"])
            if location_parts:
                loc += f" ({', '.join(location_parts)})"
            start = format_date_doc(stop["start_date"])
            end = format_date_doc(stop["end_date"])
            stop_info.append(f"{loc}: {start} - {end}")
        doc.add_paragraph(" | ".join(stop_info)).alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"Trip: {trip_name}")

    doc.add_paragraph()

    # --- Group items by day ---
    sorted_items = sorted(
        items, key=lambda x: datetime.fromisoformat(x["datetime_start"])
    )
    days = {}
    for item in sorted_items:
        date_key = datetime.fromisoformat(item["datetime_start"]).strftime("%Y-%m-%d")
        if date_key not in days:
            days[date_key] = []
        days[date_key].append(item)

    grand_total = 0
    confirmed_total = 0
    estimated_total = 0

    # Symbol for base currency
    base_symbol = get_currency_symbol(base_currency)

    # --- Process each day ---
    for date_key, day_items in days.items():
        dt_obj = datetime.strptime(date_key, "%Y-%m-%d")
        doc.add_heading(
            f"📅 {dt_obj.strftime('%A, %d-%m-%Y')}",
            level=1,
        )

        # Table: Time, Description, Type, Cost (in base), Receipt
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Time"
        hdr[1].text = "Description"
        hdr[2].text = "Type"
        hdr[3].text = f"Cost ({base_currency})"
        hdr[4].text = "Receipt"

        day_total = 0
        for item in day_items:
            dt = datetime.fromisoformat(item["datetime_start"])
            time_str = dt.strftime("%H:%M")
            orig_cost = item.get("cost", 0)
            snapshot_rate = item.get("exchange_rate_snapshot", 1.0)
            converted_cost = orig_cost * snapshot_rate
            day_total += converted_cost
            grand_total += converted_cost
            if item.get("is_confirmed"):
                confirmed_total += converted_cost
            else:
                estimated_total += converted_cost

            row_cells = table.add_row().cells
            row_cells[0].text = time_str
            row_cells[1].text = (
                f"{item['description']} ({item.get('confirmation_code', '')})"
            )
            row_cells[2].text = item["item_type"]
            # Show converted cost with base symbol; optionally include original
            display_cost = f"{base_symbol}{converted_cost:.2f}"
            orig_currency = item.get("cost_currency", "")
            if orig_currency and orig_currency != base_currency:
                display_cost += f" ({orig_cost:.2f} {orig_currency})"
            row_cells[3].text = display_cost

            # Receipt column
            receipt_path = item.get("receipt_path")
            if receipt_path and os.path.exists(receipt_path):
                try:
                    p = row_cells[4].paragraphs[0]
                    r = p.add_run()
                    r.add_picture(receipt_path, width=Inches(1.2))
                    p = row_cells[4].add_paragraph()
                    p.add_run(os.path.basename(receipt_path)).font.size = Pt(6)
                except Exception:
                    row_cells[4].text = f"📎 {os.path.basename(receipt_path)}"
            else:
                row_cells[4].text = "—"

        # Day total row
        total_row = table.add_row().cells
        total_row[0].text = ""
        total_row[1].text = ""
        total_row[2].text = "**Day Total**"
        total_row[3].text = f"{base_symbol}{day_total:.2f}"
        total_row[4].text = ""

    # --- Summary Totals ---
    doc.add_page_break()
    doc.add_heading("💰 Summary Totals", level=1)

    summary_table = doc.add_table(rows=3, cols=2)
    summary_table.style = "Light Grid Accent 1"
    summary_table.cell(0, 0).text = "Category"
    summary_table.cell(0, 1).text = "Amount"
    summary_table.cell(1, 0).text = "Total Confirmed (Booked)"
    summary_table.cell(1, 1).text = f"{base_symbol}{confirmed_total:.2f}"
    summary_table.cell(2, 0).text = "Total Estimated (Quoted)"
    summary_table.cell(2, 1).text = f"{base_symbol}{estimated_total:.2f}"

    doc.add_paragraph(f"\n**Grand Total: {base_symbol}{grand_total:.2f}**")

    if trip_budget > 0:
        remaining = trip_budget - grand_total
        doc.add_paragraph(
            f"Budget: {base_symbol}{trip_budget:.2f}  |  Remaining: {base_symbol}{remaining:.2f}"
        )

    # Footnote about conversion
    doc.add_paragraph(
        f"Note: All amounts converted to {base_currency} using the exchange rate at the time of each expense (snapshot).",
        style="Normal",
    )

    # --- Footer ---
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(
        f'Generated on {datetime.now().strftime("%d-%m-%Y at %I:%M %p")}  |  Base Currency: {base_currency}'
    )
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True

    # --- Save and return ---
    os.makedirs("generated_expense_reports", exist_ok=True)
    filename = f"generated_expense_reports/{executive['name']}_{trip_name}_ExpenseReport_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    doc.save(filename)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream, filename
